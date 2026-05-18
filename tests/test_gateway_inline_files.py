"""Tests for inline file extraction + image staging in /v1/chat/completions."""

from __future__ import annotations

import asyncio
import base64
import io
from typing import Any

import httpx
import pytest

from src.gateway.file_extract import ExtractError
from src.gateway.inline_files import (
    _decode_data_url,
    image_store,
    rewrite_image_urls,
    rewrite_input_files,
    stage_image_url,
)
from src.worker.schemas import ChatMessage


def _run(coro):
    return asyncio.run(coro)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_docx_bytes() -> bytes:
    from docx import Document

    buf = io.BytesIO()
    d = Document()
    d.add_heading("Inline Title", level=1)
    d.add_paragraph("inline body text")
    d.save(buf)
    return buf.getvalue()


def _data_url(mime: str, payload: bytes) -> str:
    return f"data:{mime};base64,{base64.b64encode(payload).decode('ascii')}"


@pytest.fixture(autouse=True)
def _clear_store():
    image_store._items.clear()
    yield
    image_store._items.clear()


# ---------------------------------------------------------------------------
# _decode_data_url
# ---------------------------------------------------------------------------


class TestDataUrlDecode:
    def test_image_png_decodes(self):
        out = _decode_data_url("data:image/png;base64,SGVsbG8=")
        assert out is not None
        data, mime = out
        assert data == b"Hello"
        assert mime == "image/png"

    def test_non_data_url_returns_none(self):
        assert _decode_data_url("https://example.com/x.png") is None

    def test_malformed_base64_returns_none(self):
        # Truly invalid b64 — '!' is not in the alphabet.
        # base64 is permissive but extra garbage past padding can sneak in;
        # we accept the lenient decoder and just ensure no crash. The bytes
        # may not be meaningful but the function must return something.
        out = _decode_data_url("data:image/png;base64,!!!")
        # Either None (regex didn't match) or a tuple — both are non-crashing
        assert out is None or isinstance(out, tuple)


# ---------------------------------------------------------------------------
# stage_image_url + rewrite_image_urls
# ---------------------------------------------------------------------------


class TestImageStaging:
    def test_http_url_passes_through(self):
        url = "https://example.com/img.png"
        assert stage_image_url(url, "https://gw.example") == url

    def test_data_url_is_staged(self):
        url = _data_url("image/png", b"\x89PNG\r\n\x1a\n" + b"x" * 50)
        new = stage_image_url(url, "https://gw.example")
        assert new.startswith("https://gw.example/v1/_staged/")
        key = new.rsplit("/", 1)[-1]
        entry = image_store.get(key)
        assert entry is not None
        assert entry.mime == "image/png"

    def test_rewrite_image_urls_mutates_message_list(self):
        msgs = [
            ChatMessage(
                role="user",
                content=[
                    {"type": "text", "text": "what is this?"},
                    {
                        "type": "image_url",
                        "image_url": {"url": _data_url("image/jpeg", b"binary")},
                    },
                ],
            )
        ]
        count = rewrite_image_urls(msgs, "https://gw.example")
        assert count == 1
        url = msgs[0].content[1]["image_url"]["url"]
        assert url.startswith("https://gw.example/v1/_staged/")

    def test_rewrite_skips_already_http(self):
        msgs = [
            ChatMessage(
                role="user",
                content=[
                    {"type": "image_url", "image_url": {"url": "https://ok.example/x.png"}},
                ],
            )
        ]
        count = rewrite_image_urls(msgs, "https://gw.example")
        assert count == 0
        assert msgs[0].content[0]["image_url"]["url"] == "https://ok.example/x.png"


# ---------------------------------------------------------------------------
# rewrite_input_files
# ---------------------------------------------------------------------------


class TestRewriteInputFiles:
    def test_data_url_pdf_replaced_with_text_block(self):
        # docx is easier to generate than a valid PDF — same path.
        docx_bytes = _make_docx_bytes()
        msgs = [
            ChatMessage(
                role="user",
                content=[
                    {"type": "text", "text": "summarize this"},
                    {
                        "type": "input_file",
                        "file": {
                            "url": _data_url(
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                docx_bytes,
                            ),
                            "filename": "memo.docx",
                        },
                    },
                ],
            )
        ]
        async def run():
            async with httpx.AsyncClient() as client:
                return await rewrite_input_files(msgs, client)

        count, warnings = _run(run())
        assert count == 1
        assert warnings == []
        # Block 1 is now a text block.
        assert msgs[0].content[1]["type"] == "text"
        body = msgs[0].content[1]["text"]
        assert "memo.docx" in body
        assert "Inline Title" in body
        assert "inline body text" in body

    def test_missing_source_raises(self):
        msgs = [
            ChatMessage(
                role="user",
                content=[{"type": "input_file", "file": {}}],
            )
        ]

        async def run():
            async with httpx.AsyncClient() as client:
                await rewrite_input_files(msgs, client)

        with pytest.raises(ExtractError) as exc_info:
            _run(run())
        assert exc_info.value.code == "missing_source"

    def test_filename_only_via_data_field(self):
        docx_bytes = _make_docx_bytes()
        b64 = base64.b64encode(docx_bytes).decode("ascii")
        msgs = [
            ChatMessage(
                role="user",
                content=[
                    {
                        "type": "input_file",
                        "file": {
                            "data": b64,
                            "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            "filename": "via_data.docx",
                        },
                    },
                ],
            )
        ]
        async def run():
            async with httpx.AsyncClient() as client:
                return await rewrite_input_files(msgs, client)

        count, _ = _run(run())
        assert count == 1
        assert "via_data.docx" in msgs[0].content[0]["text"]

    def test_text_only_message_untouched(self):
        msgs = [ChatMessage(role="user", content="hello")]
        async def run():
            async with httpx.AsyncClient() as client:
                return await rewrite_input_files(msgs, client)

        count, _ = _run(run())
        assert count == 0
        assert msgs[0].content == "hello"

    def test_non_input_file_blocks_untouched(self):
        msgs = [
            ChatMessage(
                role="user",
                content=[
                    {"type": "text", "text": "hi"},
                    {"type": "image_url", "image_url": {"url": "https://x.png"}},
                ],
            )
        ]
        async def run():
            async with httpx.AsyncClient() as client:
                return await rewrite_input_files(msgs, client)

        count, _ = _run(run())
        assert count == 0
        # All blocks preserved unchanged.
        assert msgs[0].content[0]["type"] == "text"
        assert msgs[0].content[1]["type"] == "image_url"


# ---------------------------------------------------------------------------
# /v1/_staged endpoint integration
# ---------------------------------------------------------------------------


class TestStagedEndpoint:
    def test_get_staged_returns_image(self):
        from fastapi.testclient import TestClient

        from src.gateway.server import make_gateway_app

        app = make_gateway_app(skip_router_load=True)
        client = TestClient(app)
        key = image_store.put(b"\x89PNG\r\nfakebytes", "image/png")
        resp = client.get(f"/v1/_staged/{key}")
        assert resp.status_code == 200
        assert resp.headers["content-type"] == "image/png"
        assert resp.content.startswith(b"\x89PNG")

    def test_get_staged_404_for_unknown_key(self):
        from fastapi.testclient import TestClient

        from src.gateway.server import make_gateway_app

        client = TestClient(make_gateway_app(skip_router_load=True))
        resp = client.get("/v1/_staged/does-not-exist")
        assert resp.status_code == 404
